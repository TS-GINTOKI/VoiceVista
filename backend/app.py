# backend/app.py (The Final, Complete, Working Version)
import os
import jwt
import whisper
import threading
from functools import wraps
from flask import Flask, jsonify, request, send_from_directory, Response
from flask_cors import CORS
from dotenv import load_dotenv
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from db import get_db_connection
import uuid
from datetime import datetime
import io
from docx import Document
from fpdf import FPDF

load_dotenv()
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'a-very-secret-key')

# File upload configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'wav', 'mp3', 'm4a', 'flac', 'ogg'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Create uploads directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# THE FINAL CORS FIX: Specific and correct configuration
CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}} )

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# --- JWT Auth Decorator ---
def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        if 'authorization' in request.headers:
            token = request.headers['authorization'].split(' ')[1]
        if not token:
            return jsonify({'message': 'Token is missing!'}), 401
        try:
            data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
            current_user_id = data['user_id']
        except jwt.ExpiredSignatureError:
            return jsonify({'message': 'Token has expired!'}), 401
        except jwt.InvalidTokenError:
            return jsonify({'message': 'Token is invalid!'}), 401
        except Exception as e:
            return jsonify({'message': 'Token validation failed!'}), 401
        return f(current_user_id, *args, **kwargs)
    return decorated

# --- Auth Routes ---
@app.route('/api/auth/register', methods=['POST'])
def register():
    data = request.get_json()
    hashed_password = generate_password_hash(data['password'])
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO users (name, email, password_hash) VALUES (%s, %s, %s) RETURNING id",
            (data['name'], data['email'], hashed_password)
        )
        user_id = cur.fetchone()[0]
        cur.execute("INSERT INTO user_settings (user_id) VALUES (%s)", (user_id,))
        conn.commit()
        return jsonify({'message': 'New user created!'}), 201
    except Exception as e:
        conn.rollback()
        print(f"Registration error: {str(e)}")
        return jsonify({'message': 'Registration failed', 'error': str(e)}), 500
    finally:
        cur.close()
        conn.close()

@app.route('/api/auth/login', methods=['POST'])
def login():
    auth = request.get_json()
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("SELECT id, name, email, password_hash, avatar_url FROM users WHERE email = %s", (auth['email'],))
        user = cur.fetchone()
        if not user:
            return jsonify({'message': 'Invalid email or password'}), 401
        
        if not check_password_hash(user[3], auth['password']):
            return jsonify({'message': 'Invalid email or password'}), 401
        
        token = jwt.encode({'user_id': user[0]}, app.config['SECRET_KEY'], algorithm='HS256')
        return jsonify({
            'token': token,
            'user': {'id': user[0], 'name': user[1], 'email': user[2], 'avatar_url': user[4]}
        })
    finally:
        cur.close()
        conn.close()

# --- User Profile Route ---
@app.route('/api/user/profile', methods=['GET'])
@token_required
def get_user_profile(current_user_id):
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT u.id, u.name, u.email, u.avatar_url, u.created_at,
                   us.transcription_language, us.voice_diarization, us.export_format
            FROM users u
            LEFT JOIN user_settings us ON u.id = us.user_id
            WHERE u.id = %s
        """, (current_user_id,))
        user = cur.fetchone()
        
        if not user:
            return jsonify({'message': 'User not found'}), 404
            
        return jsonify({
            'id': user[0],
            'name': user[1],
            'email': user[2],
            'avatar_url': user[3],
            'created_at': user[4].isoformat() if user[4] else None,
            'settings': {
                'transcription_language': user[5],
                'voice_diarization': user[6],
                'export_format': user[7]
            }
        })
    finally:
        cur.close()
        conn.close()

# --- Update User Profile Route ---
@app.route('/api/user/profile', methods=['PUT'])
@token_required
def update_user_profile(current_user_id):
    data = request.get_json()
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        # Update user basic info
        cur.execute("""
            UPDATE users 
            SET name = %s, email = %s
            WHERE id = %s
        """, (data.get('name'), data.get('email'), current_user_id))
        
        # Update user settings
        cur.execute("""
            UPDATE user_settings 
            SET transcription_language = %s, voice_diarization = %s, export_format = %s
            WHERE user_id = %s
        """, (
            data.get('settings', {}).get('transcription_language', 'English'),
            data.get('settings', {}).get('voice_diarization', True),
            data.get('settings', {}).get('export_format', 'pdf'),
            current_user_id
        ))
        
        conn.commit()
        return jsonify({'message': 'Profile updated successfully'})
    except Exception as e:
        conn.rollback()
        return jsonify({'message': 'Profile update failed', 'error': str(e)}), 500
    finally:
        cur.close()
        conn.close()

# --- Avatar Upload Route ---
@app.route('/api/user/avatar', methods=['POST'])
@token_required
def upload_avatar(current_user_id):
    if 'avatar' not in request.files:
        return jsonify({'message': 'No avatar file provided'}), 400
    
    file = request.files['avatar']
    if file.filename == '':
        return jsonify({'message': 'No file selected'}), 400
    
    # Check if it's an image file
    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif'}
    if not ('.' in file.filename and file.filename.rsplit('.', 1)[1].lower() in allowed_extensions):
        return jsonify({'message': 'Invalid file type. Allowed: png, jpg, jpeg, gif'}), 400
    
    # Generate unique filename
    file_extension = file.filename.rsplit('.', 1)[1].lower()
    unique_filename = f"avatar_{current_user_id}_{uuid.uuid4()}.{file_extension}"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
    
    # Save file
    file.save(file_path)
    
    # Update database
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            UPDATE users 
            SET avatar_url = %s
            WHERE id = %s
        """, (unique_filename, current_user_id))
        conn.commit()
        
        return jsonify({
            'message': 'Avatar uploaded successfully',
            'avatar_url': unique_filename
        })
    except Exception as e:
        conn.rollback()
        return jsonify({'message': 'Avatar upload failed', 'error': str(e)}), 500
    finally:
        cur.close()
        conn.close()

# --- File Upload Route ---
@app.route('/api/upload', methods=['POST'])
@token_required
def upload_audio(current_user_id):
    if 'audio' not in request.files:
        return jsonify({'message': 'No audio file provided'}), 400
    
    file = request.files['audio']
    if file.filename == '':
        return jsonify({'message': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'message': 'Invalid file type. Allowed: wav, mp3, m4a, flac, ogg'}), 400
    
    # Check file size (50MB limit)
    file.seek(0, 2)  # Seek to end
    file_size = file.tell()
    file.seek(0)  # Reset to beginning
    
    if file_size > 50 * 1024 * 1024:  # 50MB
        return jsonify({'message': 'File too large. Maximum size: 50MB'}), 400
    
    # Generate unique filename
    file_extension = file.filename.rsplit('.', 1)[1].lower()
    unique_filename = f"{uuid.uuid4()}.{file_extension}"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
    
    # Save file
    file.save(file_path)
    
    # Save to database
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            INSERT INTO audio_files (user_id, filename, original_filename, status)
            VALUES (%s, %s, %s, %s) RETURNING id
        """, (current_user_id, unique_filename, file.filename, 'processing'))
        file_id = cur.fetchone()[0]
        conn.commit()
        
        # Start transcription in background
        def process_transcription():
            try:
                print(f"Starting transcription for file_id: {file_id}")
                model = whisper.load_model("base")
                result = model.transcribe(file_path)
                transcript = result["text"]

                # Summarize
                try:
                    from transformers import pipeline
                    summarizer = pipeline('summarization', model='facebook/bart-large-cnn')
                    summary = summarizer(transcript, max_length=60, min_length=20, do_sample=False)[0]['summary_text']
                except Exception:
                    summary = transcript[:200] + ('...' if len(transcript) > 200 else '')

                # Update database with transcript and summary
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("""
                    UPDATE audio_files 
                    SET transcript = %s, summary = %s, status = 'completed'
                    WHERE id = %s
                """, (transcript, summary, file_id))
                conn.commit()
                cur.close()
                conn.close()
                print(f"Transcription completed for file_id: {file_id}")

            except Exception as e:
                print(f"Transcription failed for file_id {file_id}: {str(e)}")
                # Update status to failed
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("""
                    UPDATE audio_files 
                    SET status = 'failed'
                    WHERE id = %s
                """, (file_id,))
                conn.commit()
                cur.close()
                conn.close()
        
        # Start transcription processing in background
        thread = threading.Thread(target=process_transcription)
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'message': 'File uploaded successfully',
            'file_id': file_id,
            'filename': file.filename
        }), 201
        
    except Exception as e:
        conn.rollback()
        return jsonify({'message': 'Upload failed', 'error': str(e)}), 500
    finally:
        cur.close()
        conn.close()

# --- Transcriptions Route ---
@app.route('/api/transcriptions', methods=['GET'])
@token_required
def get_transcriptions(current_user_id):
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT id, original_filename, status, uploaded_at, transcript, summary
            FROM audio_files 
            WHERE user_id = %s 
            ORDER BY uploaded_at DESC
        """, (current_user_id,))
        transcriptions = cur.fetchall()
        result = [{
            'id': t[0], 
            'title': t[1], 
            'status': t[2], 
            'date': t[3].strftime('%Y-%m-%d %H:%M:%S') if t[3] else None,
            'transcript': t[4],
            'summary': t[5]
        } for t in transcriptions]
        return jsonify(result)
    finally:
        cur.close()
        conn.close()

# --- Delete ALL Transcriptions for current user ---
@app.route('/api/transcriptions', methods=['DELETE'])
@token_required
def delete_all_transcriptions(current_user_id):
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            DELETE FROM audio_files
            WHERE user_id = %s
        """, (current_user_id,))
        conn.commit()
        return jsonify({'message': 'All transcriptions deleted successfully'})
    except Exception as e:
        conn.rollback()
        return jsonify({'message': 'Failed to delete all transcriptions', 'error': str(e)}), 500
    finally:
        cur.close()
        conn.close()

# --- Get Single Transcription ---
@app.route('/api/transcriptions/<int:transcription_id>', methods=['GET'])
@token_required
def get_transcription(current_user_id, transcription_id):
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT id, original_filename, status, uploaded_at, transcript, summary
            FROM audio_files 
            WHERE id = %s AND user_id = %s
        """, (transcription_id, current_user_id))
        transcription = cur.fetchone()
        
        if not transcription:
            return jsonify({'message': 'Transcription not found'}), 404
            
        return jsonify({
            'id': transcription[0],
            'title': transcription[1],
            'status': transcription[2],
            'date': transcription[3].strftime('%Y-%m-%d %H:%M:%S') if transcription[3] else None,
            'transcript': transcription[4],
            'summary': transcription[5]
        })
    finally:
        cur.close()
        conn.close()

# --- Delete Single Transcription ---
@app.route('/api/transcriptions/<int:transcription_id>', methods=['DELETE'])
@token_required
def delete_transcription(current_user_id, transcription_id):
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        # Ensure it belongs to the current user
        cur.execute("""
            DELETE FROM audio_files
            WHERE id = %s AND user_id = %s
            RETURNING id
        """, (transcription_id, current_user_id))
        deleted = cur.fetchone()
        if not deleted:
            conn.rollback()
            return jsonify({'message': 'Transcription not found'}), 404
        conn.commit()
        return jsonify({'message': 'Transcription deleted successfully'})
    except Exception as e:
        conn.rollback()
        return jsonify({'message': 'Failed to delete transcription', 'error': str(e)}), 500
    finally:
        cur.close()
        conn.close()

# --- Download Transcription Route (with format and summary) ---
@app.route('/api/transcriptions/<int:transcription_id>/download', methods=['GET'])
@token_required
def download_transcription(current_user_id, transcription_id):
    try:
        from transformers import pipeline
        summarizer = pipeline('summarization', model='facebook/bart-large-cnn')
    except Exception:
        summarizer = None
    try:
        from deep_translator import GoogleTranslator
    except Exception:
        GoogleTranslator = None

    format = request.args.get('format', 'txt').lower()
    lang = request.args.get('lang', 'en')

    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT original_filename, transcript, status
            FROM audio_files 
            WHERE id = %s AND user_id = %s
        """, (transcription_id, current_user_id))
        transcription = cur.fetchone()
        
        if not transcription:
            return jsonify({'message': 'Transcription not found'}), 404
        
        if transcription[2] != 'completed':
            return jsonify({'message': 'Transcription not completed yet'}), 400
        
        transcript = transcription[1]
        summary = None
        # Summarize if possible
        if summarizer:
            try:
                summary = summarizer(transcript, max_length=60, min_length=20, do_sample=False)[0]['summary_text']
            except Exception:
                summary = None
        if not summary:
            summary = transcript[:200] + ('...' if len(transcript) > 200 else '')
        
        # Translate if needed
        lang_map = {
            'en': 'english', 'hi': 'hindi', 'ja': 'japanese', 'ko': 'korean', 'zh': 'chinese',
            'es': 'spanish', 'de': 'german', 'ru': 'russian', 'co': 'colombian spanish'
        }
        if lang != 'en' and GoogleTranslator:
            try:
                transcript = GoogleTranslator(source='auto', target=lang).translate(transcript)
                summary = GoogleTranslator(source='auto', target=lang).translate(summary)
            except Exception:
                pass
        
        filename_base = transcription[0]
        if format == 'pdf':
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, f'Transcription of: {filename_base}', ln=True)
            pdf.set_font('Arial', '', 12)
            pdf.multi_cell(0, 10, f'Summary:\n{summary}\n\nTranscript:\n{transcript}')
            pdf_bytes = pdf.output(dest='S').encode('latin1')
            return Response(pdf_bytes, mimetype='application/pdf', headers={
                'Content-Disposition': f'attachment; filename="{filename_base}_transcript.pdf"'
            })
        elif format == 'docx':
            doc = Document()
            doc.add_heading(f'Transcription of: {filename_base}', 0)
            doc.add_heading('Summary', level=1)
            doc.add_paragraph(summary)
            doc.add_heading('Transcript', level=1)
            doc.add_paragraph(transcript)
            f = io.BytesIO()
            doc.save(f)
            f.seek(0)
            return Response(f.read(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={
                'Content-Disposition': f'attachment; filename="{filename_base}_transcript.docx"'
            })
        else:
            content = f"Transcription of: {filename_base}\n\nSummary:\n{summary}\n\nTranscript:\n{transcript}"
            return Response(content, mimetype='text/plain', headers={
                'Content-Disposition': f'attachment; filename="{filename_base}_transcript.txt"'
            })
    finally:
        cur.close()
        conn.close()

# --- Serve Uploaded Files (Profile DP, Audio, etc.) ---
@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
